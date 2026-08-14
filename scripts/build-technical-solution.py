#!/usr/bin/env python3
"""Sinh tài liệu Technical Solution Phase 1 dạng .docx.

Bản rút gọn, tổng hợp từ docs/technical-solution/TS-*.md, Blueprint BA v1.11,
docs/requirements-alignment/ và kết quả chẩn đoán template hợp đồng thật.

Usage:
    pip install python-docx
    python3 scripts/build-technical-solution.py [output.docx]
"""

from __future__ import annotations

import sys
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)
DANGER = RGBColor(0xA6, 0x1B, 0x1B)
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #

def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12

    for name, size, color, before in (
        ("Heading 1", 16, ACCENT, 18),
        ("Heading 2", 13, ACCENT, 14),
        ("Heading 3", 11.5, ACCENT, 10),
    ):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr in ("begin", "PAGE", "end"):
        el = OxmlElement("w:fldChar") if instr != "PAGE" else OxmlElement("w:instrText")
        if instr == "PAGE":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    inner = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Nhấn chuột phải → Update Field để hiện mục lục."
    inner.append(t)
    fld.append(inner)
    run._r.append(fld)


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    add_rich(p, text, bold=bold, italic=italic, color=color, size=size)
    return p


def add_rich(p, text, bold=False, italic=False, color=None, size=None):
    """Hỗ trợ **đậm**, *nghiêng* và `mã`, kể cả `mã` lồng trong **đậm**."""
    import re

    def emit(chunk, is_bold, is_italic):
        for part in re.split(r"(`[^`]+`)", chunk):
            if not part:
                continue
            r = p.add_run()
            if part.startswith("`") and part.endswith("`") and len(part) > 2:
                r.text = part[1:-1]
                r.font.name = MONO_FONT
                r.font.size = Pt((size or 10.5) - 1)
            else:
                r.text = part
                r.font.size = Pt(size) if size else None
            r.bold = is_bold
            r.italic = is_italic
            if color is not None:
                r.font.color.rgb = color

    # Tách **đậm** trước (lazy để không nuốt cả câu), rồi *nghiêng*, rồi `mã`.
    for seg in re.split(r"(\*\*.+?\*\*)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**") and len(seg) > 4:
            emit(seg[2:-2], True, italic)
            continue
        for sub in re.split(r"(?<!\*)(\*[^*\n]+\*)(?!\*)", seg):
            if not sub:
                continue
            if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                emit(sub[1:-1], bold, True)
            else:
                emit(sub, bold, italic)
    return p


def bullets(doc, items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        add_rich(p, it)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = MONO_FONT
    r.font.size = Pt(8.5)
    shade(p, "F4F6F8")
    return p


def shade(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    pPr.append(sh)


def callout(doc, text, color=ACCENT, fill="EAF1F8"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    add_rich(p, text, color=color)
    shade(p, fill)
    return p


def table(doc, headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_rich(p, htxt, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=font_size)
        for r in p.runs:
            r.bold = True
            if r.font.size is None:
                r.font.size = Pt(font_size)
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), "1F4E79")
        tcPr.append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            add_rich(p, str(v), size=font_size)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(font_size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# --------------------------------------------------------------------------- #
# Document
# --------------------------------------------------------------------------- #

def build(path: str) -> None:
    doc = Document()
    setup_styles(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(1.8)
    add_page_number_footer(doc)

    # ---------------- Trang bìa ----------------
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TÀI LIỆU GIẢI PHÁP KỸ THUẬT")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = ACCENT

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HỆ THỐNG AI REVIEW HỢP ĐỒNG (AI LEGAL)")
    r.bold = True
    r.font.size = Pt(15)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Phase 1 — Hợp đồng khung  ·  Saint-Gobain Việt Nam")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"Bản rút gọn 1.0 — {date.today().strftime('%d/%m/%Y')}")
    r.font.size = Pt(11)
    r.italic = True

    for _ in range(3):
        doc.add_paragraph()
    table(doc, ["Hạng mục", "Nội dung"], [
        ["Phạm vi", "Phase 1 — hợp đồng khung có template khoá cấu trúc"],
        ["Nguồn đầu vào", "Blueprint BA v1.11 · docs/requirements-alignment/ · chẩn đoán template thật"],
        ["Chủ trì", "Backend / AI Engineer"],
        ["Trạng thái", "Draft — chờ chốt PoC cổng chặn và phương án tiến độ mục XV"],
        ["Tài liệu chi tiết", "docs/technical-solution/TS-00 … TS-03 (bản đầy đủ)"],
    ], widths=[4.0, 12.5])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h1(doc, "MỤC LỤC")
    add_toc(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------- I ----------------
    h1(doc, "I. MỤC TIÊU VÀ PHẠM VI")

    para(doc, "Hệ thống AI Legal giúp phòng Mua hàng tự rà soát hợp đồng .docx bằng AI trước khi "
              "trình Purchasing Manager và Legal phê duyệt, sau đó đẩy sang FPT.eContract để trình ký. "
              "Toàn bộ suy luận AI chạy trên hạ tầng nội bộ.")

    h2(doc, "1. Phân kỳ")
    table(doc, ["Giai đoạn", "Đối tượng", "Đặc điểm kỹ thuật"], [
        ["**Phase 1**", "Hợp đồng khung (framework)",
         "Template do Legal ban hành, khoá cấu trúc bằng Restrict Editing. Chỉ các vùng Legal mở mới được sửa; phần còn lại tuyệt đối không đổi."],
        ["Phase 2", "Hợp đồng nhà cung cấp",
         "File tuỳ ý, không bắt buộc khớp template, có thể không có vùng mở nào. Redlining văn bản tự do."],
    ], widths=[2.6, 4.2, 9.7])

    h2(doc, "2. Ràng buộc cứng")
    table(doc, ["Mã", "Ràng buộc"], [
        ["C-1", "LLM chạy local, **tuyệt đối không gọi cloud AI**. Dữ liệu hợp đồng không rời hạ tầng nội bộ."],
        ["C-2", "Output `.docx` giữ format giống hệt input."],
        ["C-3", "**Không bao giờ ghi vào vùng khoá.** Allow-list Lớp 1 nằm ở tầng ghi file backend."],
        ["C-4", "PT3 reupload phát hiện vùng khoá bị sửa → chặn hoàn toàn, không có override."],
        ["C-5", "Chỉ gọi eContract sau khi Legal approve. Chiều nhận file đã ký ngoài phạm vi."],
        ["C-6", "Approval Matrix chỉ dùng cảnh báo và tính điểm, **không auto-routing**."],
        ["C-7", "Queue xử lý FIFO."],
        ["C-8", "Marker eContract chèn bằng mực trắng, id duy nhất toàn file."],
        ["C-9", "Kết quả AI chỉ là gợi ý; trách nhiệm cuối thuộc người phê duyệt."],
        ["C-10", "Checklist do Legal tự vận hành trên UI, không cần deploy."],
        ["C-11", "System Prompt quản lý bằng Git, CI validate bắt buộc pass."],
        ["C-12", "Không hardcode nội dung pháp lý trong prompt hoặc code."],
    ], widths=[1.4, 15.1])

    # ---------------- II ----------------
    h1(doc, "II. KẾT QUẢ KHẢO SÁT TEMPLATE THẬT")

    callout(doc, "Toàn bộ thiết kế dưới đây dựa trên số liệu đo được từ hợp đồng khung thật đang lưu hành, "
                 "không phải giả định. Công cụ chẩn đoán nằm trong repo: `scripts/inspect-template.py` "
                 "và `scripts/map-locked-regions.py`, chạy được trên mọi file `.docx`.")

    h2(doc, "1. Bảng phát hiện")
    table(doc, ["Mã", "Phát hiện", "Số liệu", "Hệ quả thiết kế"], [
        ["PH-1", "Cơ chế vùng mở là **Range Permission**",
         "16 cặp `w:permStart`/`w:permEnd`, `edGrp=everyone`. `w:sdt` = 0. Form Field = 0",
         "Giả định ban đầu đúng, nhưng đảo ngược khuyến nghị Word engine (xem QĐ-1)"],
        ["PH-2", "Không có Content Control nào", "0 `w:sdt` toàn tài liệu",
         "Cơ chế lock-mode `LOCK_VIOLATION` của editor thương mại **không kích hoạt** trên template này"],
        ["PH-3", "`w14:paraId` phủ toàn bộ", "230/230 đoạn có `paraId`, 0 trùng",
         "Giải được anchor cho comment và cho marker kéo-thả"],
        ["PH-4", "Vùng mở **không đồng nghĩa với field**",
         "11 atomic field (≤ 91 ký tự) · 1 vùng rỗng · 2 block đa đoạn (lớn nhất **3.174 ký tự / 19 đoạn**) · 2 vùng bắc qua ranh giới bảng",
         "Phải có **hai chế độ write-back** ngay từ Phase 1 (QĐ-3)"],
        ["PH-5", "Một đoạn có thể chứa nhiều vùng mở",
         "Đoạn 66 chứa cả `1419390840` (\"30\") và `482367384` (\"ký hợp đồng\")",
         "Khoá anchor phải là tổ hợp `(paraId, permId)`, không dùng `paraId` đơn lẻ"],
        ["PH-6", "Số điều khoản **không nằm trong text**",
         "`numbering.xml`: `lvl0 = \"Điều %1.\"`, `lvl1 = \"%1.%2\"`; Heading1–4 đều trỏ `numId=3`",
         "Bắt buộc có bộ resolve numbering, nếu không sẽ mất khả năng trích dẫn điều khoản"],
        ["PH-7", "File đã mang sẵn comment thật",
         "3 comment, tác giả gồm cả người ngoài công ty, neo trong vùng mở",
         "PA-B phải **merge**, không ghi đè; tránh đụng `w:id` comment sẵn có"],
        ["PH-8", "Restrict Editing bằng hash mật khẩu",
         "`edit=readOnly`, `enforcement=1`, `rsaAES`, có `hash` + `salt`",
         "Backend **không cần** mật khẩu để ghi — càng khẳng định allow-list Lớp 1 là bắt buộc"],
        ["PH-9", "Có lỗi dữ liệu thật trong hợp đồng",
         "`685.000.000` viết chữ 2 kiểu: \"…tám lăm **triệu**…\" và \"…tám mươi lăm **nghìn**…\"",
         "Bằng chứng cho tầng kiểm tra nhất quán deterministic, không cần LLM"],
        ["PH-10", "Bản đồ khoá/mở rất nhất quán",
         "**78,0% khoá / 22,0% mở**. Điều 6→14 khoá 100%. Điều 4 Thanh toán mở 99%. 10 đoạn hỗn hợp",
         "Write-back phải ở **cấp run**, không phải cấp paragraph"],
    ], widths=[1.2, 3.4, 5.4, 6.5], font_size=8.5)

    h2(doc, "2. Bản đồ khoá / mở theo điều khoản")
    para(doc, "Legal khoá toàn bộ khung pháp lý bảo vệ và chỉ mở các thông số thương mại của từng thương vụ.")
    table(doc, ["Điều", "Tên", "% mở", "Điều", "Tên", "% mở"], [
        ["1", "Hàng Hóa", "17%", "8", "An toàn lao động", "0%"],
        ["2", "Đơn Đặt Hàng", "0%", "9", "Phạt vi phạm & bồi thường", "0%"],
        ["3", "Giao hàng", "8%", "10", "Sự kiện bất khả kháng", "0%"],
        ["**4**", "**Thanh toán**", "**99%**", "11", "Bảo mật thông tin", "0%"],
        ["**5**", "**Bảo hành**", "**96%**", "12", "Chống tham nhũng", "0%"],
        ["6", "Cam kết về Hàng Hóa", "0%", "13", "Thời hạn & chấm dứt", "0%"],
        ["7", "Cam kết năng lực Bên Bán", "0%", "14", "Điều khoản cuối cùng", "0%"],
    ], widths=[1.3, 4.8, 1.5, 1.3, 4.8, 1.5])
    para(doc, "Khối thông tin các bên, khối chữ ký và Phụ lục 01 nằm ngoài đánh số điều, đều là vùng mở.")

    h2(doc, "3. Ba điều chỉnh tiền đề")
    bullets(doc, [
        "**Phase 1 không thuần là bài toán điền field.** Đúng với 11/16 vùng, nhưng vùng Thanh toán dài 19 đoạn "
        "và Phụ lục 01 gồm cả bảng là redlining văn bản tự do — bài toán vốn xếp cho Phase 2 nhưng đã xuất hiện ngay ở Phase 1.",
        "**Vùng mở nằm giữa câu bị khoá.** Ví dụ điều 3.1: *\"Bên Bán giao hàng trong vòng [30] ngày kể từ ngày "
        "[ký hợp đồng]\"* — khung câu do Legal khoá, chỉ hai giá trị được mở. Đây là ràng buộc cứng cho tầng ghi.",
        "**Nhu cầu sửa vùng khoá là có thật.** Comment thật trong file đề nghị chỉnh điều 3.5 và 3.6, mà hai điều này "
        "khoá 100%. Blueprint hiện chỉ trả lời \"Loại B — chỉ annotation\", chưa có đường escalate. Cần BA bổ sung.",
    ])

    # ---------------- III ----------------
    h1(doc, "III. CÁC QUYẾT ĐỊNH KIẾN TRÚC")

    table(doc, ["Mã", "Quyết định", "Lý do cốt lõi"], [
        ["**QĐ-1**",
         "Backend sở hữu sự thật OOXML bằng **Python + `lxml`**. Không giao việc ghi cho engine bên thứ ba.",
         "C-3 đòi hỏi chứng minh được không byte nào ngoài vùng mở bị đổi — chỉ làm được khi kiểm soát toàn bộ đường ghi. "
         "PH-2 đã loại bỏ lý do chính để chọn editor thương mại cho đường ghi. Phạm vi ghi thật sự hẹp: thay text trong "
         "`permStart…permEnd`, chèn run mực trắng, thêm `w:comment`."],
        ["**QĐ-2**",
         "Tách **engine ghi** khỏi **engine hiển thị**. Pilot Phase 1 chạy bằng preview read-only + form trường mở.",
         "Không để việc chọn editor chặn tiến độ. TH2 Track Changes là hạng mục duy nhất thật sự cần editor nhúng, "
         "cũng là hạng mục ít dùng nhất — hoãn sang Phase 1.5."],
        ["**QĐ-3**",
         "**Hai chế độ write-back**, phân loại tự động theo hình dạng vùng mở.",
         "Hệ quả trực tiếp của PH-4: không thể dùng một đường ghi chung cho ô 2 ký tự và khối 3.174 ký tự."],
        ["**QĐ-4**",
         "Anchor model: khoá tổ hợp `(paraId, permId, ordinal)` + fingerprint dự phòng.",
         "PH-3 cho `paraId` ổn định; PH-5 chứng minh `paraId` một mình không đủ."],
        ["**QĐ-5**",
         "**DB là nguồn sự thật** cho comment và diff (PA-A luôn bật). Ghi `w:comment` vào `.docx` chỉ là export (PA-B).",
         "PH-7: file đã mang comment của bên thứ ba; nếu coi file là nguồn sự thật thì mỗi vòng PT3 sẽ nuốt mất hoặc nhân bản comment."],
        ["**QĐ-6**",
         "Điểm số do **code tính**, LLM không sinh số.",
         "Hệ thống pháp chế: mỗi con số phải giải thích được và tái lập được trước Legal và audit."],
        ["**QĐ-7**",
         "Blob file ở **MinIO**, DB chỉ giữ metadata + hash.",
         "MinIO self-hosted vẫn nằm trọn trong hạ tầng nội bộ nên không vi phạm NFR-S1/S5, đồng thời tránh phình DB. "
         "Đây là phản biện quyết định D3 của PM."],
        ["**QĐ-8**",
         "**Structural binding bắt buộc**; đường chính là instantiate-from-template.",
         "Vá lỗ hổng của Blueprint v1.6: bỏ so khớp *nội dung* là đúng, nhưng bỏ luôn ràng buộc *cấu trúc* thì mô hình "
         "bảo vệ vùng khoá sụp đổ."],
    ], widths=[1.6, 5.4, 9.5], font_size=9)

    # ---------------- IV ----------------
    h1(doc, "IV. KIẾN TRÚC TỔNG THỂ")

    h2(doc, "1. Công nghệ")
    table(doc, ["Lớp", "Lựa chọn", "Ghi chú"], [
        ["API", "Python 3.12 + FastAPI", "Cùng ngôn ngữ với tầng AI; Pydantic sinh OpenAPI khớp type FE"],
        ["Worker", "Celery + Redis", "Job AI chạy dài, cần persistent queue"],
        ["CSDL", "PostgreSQL 16", "JSONB cho findings/intake, transactional outbox"],
        ["Lưu trữ file", "MinIO (S3-compatible, self-hosted)", "Mọi version `.docx`, mã hoá at-rest"],
        ["LLM serving", "vLLM, endpoint OpenAI-compatible", "Batching + guided decoding"],
        ["LLM chính", "Qwen3-30B-A3B (MoE, Apache 2.0)", "FP8 trên A100 80GB; AWQ 4-bit nếu là 40GB"],
        ["Embedding / Rerank", "BGE-M3 + `bge-reranker-v2-m3`", "Khớp clause ↔ đoạn văn bản"],
        ["OOXML", "`lxml` (thư viện nội bộ `services/document/`)", "QĐ-1"],
        ["Frontend", "Next.js 14 + TypeScript + Tailwind + shadcn/ui", "Giữ nguyên nền demo"],
        ["Triển khai", "Docker Compose (Sprint 1)", "Lên K8s khi cần"],
    ], widths=[3.0, 5.6, 7.9])

    h2(doc, "2. Phân rã module backend")
    code(doc,
         "app/\n"
         "├── api/          routers mỏng: auth, reviews, documents, comments, markers, files, events(SSE), config/, callbacks\n"
         "├── domain/       entities · enums · state_machine · rbac · errors\n"
         "├── services/\n"
         "│   ├── document/ ★ engine · lxml_engine · ooxml_reader · region_classifier\n"
         "│   │              writer_inline · writer_block · allowlist(Lớp 1) · anchor\n"
         "│   │              comment · marker · numbering_resolver · structural_binding\n"
         "│   ├── ai/       ★ pipeline · segmenter · matcher · judge · aggregator\n"
         "│   │              scorer · narrator · fallback · injection_guard · schemas/\n"
         "│   └── review/ · config/ · econtract/ · identity/ · storage/\n"
         "├── workers/      ai_review · ai_chat · field_validation · econtract_push\n"
         "│                 econtract_reconcile · retention · eval_golden\n"
         "├── infra/        db · redis · minio · vllm_client · embed_client · settings · observability\n"
         "└── prompts/      loader đọc /prompts từ Git, render placeholder, cache")
    para(doc, "Hai module đánh dấu ★ tập trung toàn bộ độ khó. Chúng được thiết kế như **thư viện độc lập, "
              "test được riêng**, không phụ thuộc FastAPI.")

    h2(doc, "3. Phân bổ GPU và capacity")
    para(doc, "Hạ tầng: **2 × NVIDIA A100**. Card 0 chạy LLM chính, card 1 chạy embedding + reranker và làm "
              "headroom cho burst — cách ly lỗi, deploy/rollback model độc lập.")
    table(doc, ["Dung lượng card", "Model", "Lượng tử", "`max_model_len`", "`max_num_seqs`"], [
        ["**80 GB** (mặc định thiết kế)", "Qwen3-30B-A3B", "FP8", "32768", "64"],
        ["40 GB (dự phòng)", "Qwen3-30B-A3B", "AWQ 4-bit", "16384", "32"],
        ["40 GB, nếu AWQ tụt chất lượng", "Qwen3-14B", "FP8", "16384", "48"],
    ], widths=[5.2, 3.6, 2.6, 2.8, 2.3])

    para(doc, "Ước tính cho một hợp đồng, giả định checklist có 40 clause:")
    code(doc,
         "41 lần gọi LLM  ≈ 103.000 token vào · 14.600 token ra\n"
         "prefill  ~15 s   decode  ~10 s   embedding & rerank  ~10 s\n"
         "OOXML parse & write-back  ~8 s    DB / hàng đợi / overhead  ~15 s\n"
         "────────────────────────────────────────────────────────────────\n"
         "tổng ≈ 58 s / hợp đồng   ⇒  NFR-P2 (≤ 10 phút) đạt với biên rất rộng")
    para(doc, "Tài liệu thật chỉ 23.313 ký tự (~8–10K token) nên **hiệu năng không phải vấn đề của dự án này**. "
              "Ưu tiên độ chính xác và khả năng truy vết, không tối ưu throughput sớm.")

    # ---------------- V ----------------
    h1(doc, "V. TẦNG TÀI LIỆU — OOXML VÀ BẢO VỆ VÙNG KHOÁ")

    callout(doc, "Đây là phần rủi ro cao nhất của dự án và là cổng chặn kỹ thuật. Mọi thiết kế trong mục này "
                 "phải chứng minh được: **không một ký tự nào trong vùng khoá bị thay đổi — kể cả khi LLM bị lừa, "
                 "kể cả khi frontend bị bypass, kể cả khi người dùng cố tình.**", color=DANGER, fill="FBECEC")

    h2(doc, "1. Ba chế độ write-back (QĐ-3)")
    table(doc, ["Chế độ", "Áp dụng cho", "Cách ghi", "Rủi ro format"], [
        ["**A — Inline replace**", "Atomic field: `permStart`/`permEnd` trong cùng một đoạn (11/16 vùng)",
         "Thay text của các `w:r` nằm giữa hai mốc, **kế thừa `w:rPr` của run đầu tiên**, xoá run thừa", "Gần bằng 0"],
        ["**B — Block replace**", "Vùng đa đoạn hoặc trong bảng (4/16 vùng)",
         "Thay nội dung từng `w:p` trong khoảng, **giữ nguyên `w:pPr`, `w:numPr`, cấu trúc `w:tbl`**. "
         "Phase 1 không cho thêm/bớt số đoạn", "Trung bình — cần PoC"],
        ["**C — Chỉ cảnh báo**", "Vùng rỗng (không có run để kế thừa định dạng); vùng bắc qua bảng nếu PoC thất bại",
         "Không ghi. AI chỉ sinh annotation", "Không có"],
    ], widths=[3.0, 4.2, 6.8, 2.5])
    para(doc, "Phân loại chạy tự động tại bước ingestion và lưu vào `document_fields.region_kind`.")

    h2(doc, "2. Allow-list Lớp 1 — điểm enforce C-3")
    bullets(doc, [
        "Hàm ghi **chỉ nhận `(permId, new_value)`**. Không tồn tại API nào nhận toàn văn bản tài liệu.",
        "`permId` không nằm trong inventory vùng mở của đúng version tài liệu đó → từ chối, ghi audit.",
        "Endpoint `PATCH /reviews/{id}/document {text}` của bản demo **bị loại bỏ** vì phá vỡ mô hình này.",
        "Sau mỗi lần ghi, chạy kiểm chứng tự động: **diff XML ngoài allow-list phải rỗng**. Không rỗng → rollback.",
        "Chat AI (PT1): nếu yêu cầu của người dùng nhắm vào anchor ngoài allow-list thì **từ chối trước khi gọi LLM**.",
    ])

    h2(doc, "3. Anchor model (QĐ-4)")
    table(doc, ["Đối tượng neo", "Khoá chính", "Khoá dự phòng"], [
        ["Vùng mở", "`permId`", "`sha256(text đã normalize)` + thứ tự xuất hiện"],
        ["Đoạn văn (comment, marker)", "`paraId`", "`sha256(text đoạn đã normalize)` + ordinal"],
        ["Vị trí trong đoạn", "offset ký tự sau normalize", "—"],
    ], widths=[4.5, 4.5, 7.5])
    para(doc, "Khi cả hai khoá đều thất bại, comment chuyển trạng thái `orphaned` — vẫn hiển thị nhưng không neo.")

    h2(doc, "4. Resolve numbering")
    para(doc, "Vì PH-6, phải khôi phục số điều khoản bằng cách đọc `styles.xml` → `numPr` của style → "
              "`numbering.xml` → `abstractNum` → `lvlText`, rồi đếm tuần tự theo cấp. "
              "Bản prototype đã chạy đúng trên template thật (`scripts/map-locked-regions.py --outline`), "
              "khôi phục đầy đủ Điều 1–14 cùng các cấp `1.1`, `a.`, `(i)`.")

    h2(doc, "5. Structural binding (QĐ-8)")
    table(doc, ["Đường vào", "Cơ chế", "Khi không khớp"], [
        ["**Instantiate** (đường chính)", "Hệ thống sinh tài liệu từ template đã đăng ký — inventory vùng mở tin cậy tuyệt đối", "Không xảy ra"],
        ["**Upload** (đường phụ)", "Đối chiếu tập `permId` + số lượng vùng + hash nội dung vùng khoá với template đã đăng ký",
         "Chặn, trả danh sách `FieldStructureIssue[]` chỉ rõ sai ở đâu. **Không có override** (C-4)"],
    ], widths=[4.2, 8.3, 4.0])
    para(doc, "Vì `permId` không mang tên nghiệp vụ (PH-2), cần bảng `template_field_map` do Legal/IT khai báo: "
              "`permId` → tên nghiệp vụ, kiểu dữ liệu, clause liên kết, ràng buộc validate.")

    h2(doc, "6. Chèn marker ký số")
    para(doc, "Kéo-thả trên preview trả về **`paraId` + offset**, không phải toạ độ pixel. Backend chèn một `w:r` "
              "mang `w:color w:val=\"FFFFFF\"` chứa chuỗi `#ds:id r:p_001_r_001 h:100 #` tại đúng anchor đó. "
              "Marker validate hai lớp: client theo bảng lỗi FPT, server bắt buộc validate lại với cùng bộ mã lỗi.")

    # ---------------- VI ----------------
    h1(doc, "VI. TẦNG AI")

    h2(doc, "1. Pipeline")
    code(doc,
         "Stage 0    Ingestion — parse OOXML, inventory vùng mở/khoá, resolve numbering,\n"
         "                       phân đoạn thành clause unit kèm anchor\n"
         "Stage 0.5  Consistency rules  (KHÔNG dùng LLM)\n"
         "                       số ↔ chữ · giá trị HĐ ↔ tổng phụ lục · thứ tự mốc ngày\n"
         "                       tên & MST các bên · đơn vị tiền tệ · field bắt buộc còn rỗng\n"
         "Stage 1    Clause matching  (KHÔNG dùng LLM)\n"
         "                       rule-based (keywords + patterns của Legal) + semantic (embedding + rerank)\n"
         "Stage 2    Per-clause judgment  (LLM, stage checklist_review)\n"
         "                       mỗi cặp = 1 lần gọi ngắn, song song, guided JSON\n"
         "                       verdict ∈ {ideal_met, fallback_met, below_fallback,\n"
         "                                  red_line_violation, missing, not_applicable}\n"
         "Stage 3    Aggregation  (KHÔNG dùng LLM)\n"
         "                       (kind × severity × verdict) → red_flag | warning | protection | missing_protection\n"
         "                       Loại A nếu anchor trong vùng mở, Loại B nếu vùng khoá\n"
         "Stage 4    Scoring  (KHÔNG dùng LLM)\n"
         "Stage 5    Narrative  (LLM, stage ai_summary_fairness) — chỉ viết văn, không sinh số\n"
         "Stage 6    Write-back — proposals → allow-list → ghi OOXML → version mới")

    para(doc, "Chia theo clause thay vì nhồi cả tài liệu vào một prompt: mỗi finding truy vết được về đúng một clause "
              "và một lần gọi LLM; prompt ngắn nên chính xác hơn; lỗi cục bộ không phá cả kết quả.")

    h2(doc, "2. Bốn stage prompt")
    table(doc, ["Stage", "Khi chạy", "Đầu ra"], [
        ["`checklist_review`", "Lượt review đầu tiên trong queue", "Findings + đề xuất theo từng clause"],
        ["`chat_edit`", "PT1 — người dùng gõ yêu cầu trong chat", "Diff đề xuất trên vùng mở"],
        ["`ai_summary_fairness`", "Sau khi có findings", "Đoạn tóm tắt tiếng Việt"],
        ["`field_validation`", "Sau khi lưu một trường mở", "Đánh giá lại clause liên quan"],
    ], widths=[3.8, 6.2, 6.5])
    para(doc, "Prompt quản lý bằng file trong Git, `_shared/injection_guard.md` prepend vào mọi stage, con trỏ version "
              "qua `current.json`, CI `validate-prompts` chặn placeholder lạ và chặn hardcode nội dung pháp lý. "
              "Nội dung pháp lý luôn đến từ DB checklist qua `{{checklist_items}}`.")

    h2(doc, "3. Structured output")
    bullets(doc, [
        "**Guided decoding** của vLLM (xgrammar/outlines) với JSON Schema từng stage — không parse JSON bằng regex.",
        "Validate bằng Pydantic, retry có giới hạn.",
        "`temperature=0` và cố định `seed` cho các stage phán xét; chỉ stage tóm tắt dùng temperature > 0.",
    ])

    h2(doc, "4. Hai điểm số (QĐ-6)")
    table(doc, ["Chỉ số", "Ý nghĩa", "Đầu vào"], [
        ["**AI Confidence**", "Độ chắc chắn của bản thân phân tích. Điểm thấp nghĩa là *AI không dám chắc*, "
                              "không phải *hợp đồng xấu*",
         "Tỷ lệ clause tìm được segment khớp · độ đồng thuận giữa tầng rule-based và semantic · điểm similarity · "
         "self-confidence đã hiệu chuẩn · có phải kết quả fallback không"],
        ["**Fairness**", "Mức cân bằng của điều khoản cho Công ty",
         "Tương quan Red Flag / Missing Protection / Warning với Protection, **có trọng số theo `severity`** "
         "trong cấu hình Legal — không dùng hằng số hardcode"],
    ], widths=[3.2, 5.3, 8.0])
    para(doc, "Hai chỉ số tách biệt hoàn toàn. Trọng số điều chỉnh được qua cấu hình, không phải sửa code. "
              "Toàn bộ heuristic của `contract-insight.ts` trong bản demo bị bỏ.")

    h2(doc, "5. Khả năng tái lập và fallback")
    bullets(doc, [
        "Mỗi lần chạy ghi một bản ghi `ai_runs` đủ để **tái lập kết quả**: `model_id`, `model_hash`, `prompt_stage`, "
        "`prompt_version` (hash file), `checklist_config_version`, `document_version`, `temperature`, `seed`, token, latency.",
        "LLM lỗi hoặc timeout → chạy tầng rule-based đơn thuần, đánh dấu `is_fallback = true`, hạ AI Confidence theo "
        "quy tắc rõ ràng, hiện banner cảnh báo. Fallback **không** sinh đề xuất thay thế văn bản, chỉ cảnh báo.",
        "Chống prompt injection ba lớp: guard trong prompt · detector chạy trước, phát hiện thì gắn Red Flag và tiếp tục · "
        "và quan trọng nhất là **allow-list Lớp 1 ở tầng ghi** — dù LLM có bị lừa thì vùng khoá vẫn an toàn.",
    ])

    # ---------------- VII ----------------
    h1(doc, "VII. DỮ LIỆU, API VÀ QUY TRÌNH")

    h2(doc, "1. Nhóm bảng")
    table(doc, ["Nhóm", "Bảng", "Đặc tính"], [
        ["Định danh", "`users`, `sessions`", "`line_manager_id` tự tham chiếu"],
        ["Vòng đời review", "`contract_reviews`, `review_versions`, `review_files`", "`review_versions` **immutable**"],
        ["Tài liệu", "`document_fields`, `document_paragraphs`, `document_segments`", "`document_fields` **chính là allow-list**"],
        ["AI", "`ai_runs`, `ai_findings`, `ai_proposals`, `chat_messages`", "`ai_runs` đủ để tái lập"],
        ["Cộng tác", "`comments`, `comment_replies`, `legal_edits`, `feedback_items`, `feedback_attachments`", "Anchor theo QĐ-4"],
        ["Ký số", "`sign_recipients`, `markers`, `econtract_envelopes`, `econtract_outbox`", "Outbox pattern"],
        ["Cấu hình", "`contract_type_configs`, `contract_templates`, `checklist_clauses`, `approval_matrices`, `form_lists`", "Legal tự vận hành (C-10)"],
        ["Kiểm toán", "`audit_log`", "**Append-only**, trigger chặn UPDATE/DELETE"],
    ], widths=[2.6, 7.4, 6.5])

    h2(doc, "2. State machine")
    code(doc,
         "draft → queued → processing → reviewed → awaiting_markers\n"
         "      → pending_manager ─approve→ pending_legal ─approve→ syncing_econtract → signed\n"
         "            └─reject──┐              └─reject──┐\n"
         "                      └──────→ rejected ←──────┘  → (sửa, bump version) → submit lại\n"
         "bổ sung: ai_failed · econtract_failed · cancelled")
    para(doc, "**Backend là nơi duy nhất được chuyển trạng thái**; frontend chỉ đọc. Quy tắc chặn: không submit khi "
              "marker chưa hợp lệ, khi còn thay đổi chưa lưu, hoặc khi đang có job AI chạy. Ghi có optimistic locking "
              "bằng version/ETag.")

    h2(doc, "3. Nhóm API")
    table(doc, ["Nhóm", "Endpoint tiêu biểu", "Ghi chú"], [
        ["Xác thực & danh mục", "`/api/auth/login`, `/api/me`, `/api/form-lists/{listKey}`", "Gộp 4 endpoint danh mục rời rạc của demo"],
        ["Vòng đời review", "`/api/reviews`, `/{id}/submit`, `/{id}/decision`, `/{id}/versions`", "`decision` dùng chung cho Manager và Legal"],
        ["Tài liệu", "`GET/PUT /{id}/fields`, `/{id}/paragraphs`, `/{id}/reupload`, `/{id}/chat`", "**Ghi theo `permId`** — điểm enforce C-3"],
        ["Cộng tác", "`/{id}/comments`, `/{id}/legal-edits`", "`legal-edits` thuộc Phase 1.5"],
        ["Ký số", "`/{id}/recipients`, `/{id}/markers`, `/{id}/markers/validate`", "Validate lại server-side"],
        ["Cấu hình", "`/api/config/contract-types|clauses|matrices|templates|form-lists|audit`", "Nhóm thiết kế mới hoàn toàn"],
        ["Hệ thống", "`/api/system-prompts`, `/api/events` (SSE), `/api/callbacks/econtract`", "SSE thay polling giả lập"],
    ], widths=[3.0, 7.4, 6.1])

    h2(doc, "4. Tích hợp FPT.eContract (outbound)")
    bullets(doc, [
        "**Transactional outbox**: Legal approve commit DB xong mới đẩy job gọi FPT — tránh mất hoặc trùng.",
        "**Idempotency** theo `refId` = `review.code`. Token cache và refresh trước `expTime`.",
        "**Retry có backoff** cộng **job đối soát định kỳ** bằng API lấy link ký, phòng callback treo.",
        "Map mã lỗi FPT sang thông báo tiếng Việt hiển thị lên UI.",
        "Chưa có credentials môi trường Demo → xây **adapter + mock server** để không bị chặn tiến độ.",
    ])

    h2(doc, "5. Bảo mật")
    table(doc, ["Hạng mục", "Yêu cầu"], [
        ["RBAC", "Enforce **server-side**. Purchasing chỉ thấy review của mình; Manager thấy của cấp dưới theo Line Manager"],
        ["Truy cập file", "Presigned URL TTL ngắn qua API kiểm quyền. Không có public path"],
        ["Mã hoá", "File mã hoá at-rest; secret của FPT lưu trong secret manager, không nằm trong repo"],
        ["Audit", "Append-only, lưu **giá trị cũ → mới** cho mọi thay đổi; bắt buộc ghi cho hành động AI, sửa field, quyết định duyệt, sửa cấu hình, sửa prompt"],
        ["Dữ liệu AI", "Không có đường ra Internet từ tầng inference (C-1)"],
    ], widths=[3.0, 13.5])

    # ---------------- VIII ----------------
    h1(doc, "VIII. FRONTEND")
    table(doc, ["Việc", "Trạng thái hiện tại", "Cần làm"], [
        ["Kết nối API thật", "Mock qua `localStorage`", "Bật `NEXT_PUBLIC_USE_MOCK=false`, gỡ toàn bộ mock"],
        ["Hiển thị tài liệu", "`docx-preview`", "Preview read-only + form trường mở (QĐ-2)"],
        ["Lưu", "Autosave mỗi thao tác", "**Lưu thủ công**, cảnh báo khi thoát còn thay đổi chưa lưu"],
        ["Marker", "Click chọn vị trí định sẵn", "**Kéo-thả trên preview**, trả về `paraId` + offset"],
        ["PT2 / PT3", "Có service, chưa có UI", "Form trường mở cho PT2; nút reupload cho PT3"],
        ["TH1 comment", "Chỉ 1 comment tổng khi reject", "Thread 2 chiều neo theo đoạn"],
        ["Queue", "Polling giả lập", "SSE"],
        ["Dọn dẹp", "Còn role `legal_lead`", "Gỡ khỏi type và luồng; thêm optimistic locking"],
    ], widths=[3.4, 5.6, 7.5])

    # ---------------- IX ----------------
    h1(doc, "IX. KIỂM THỬ VÀ NGHIỆM THU")

    table(doc, ["Nhóm", "Nội dung", "Tiêu chí đạt"], [
        ["**Bảo vệ vùng khoá**", "Sau mỗi thao tác ghi, so sánh XML trước/sau",
         "Diff ngoài allow-list = rỗng. **Đây là tiêu chí không được phép trượt.**"],
        ["**Giữ format**", "Round-trip trên template thật, mở bằng Microsoft Word",
         "Word không cảnh báo file lỗi; số trang, bảng, numbering không đổi"],
        ["**Chất lượng AI**", "Golden set do Legal gán nhãn, đo theo từng clause",
         "Ưu tiên **recall của nhóm Red Flag / Block** — thà báo thừa còn hơn bỏ sót. Ngưỡng cụ thể chờ PM chốt"],
        ["**Regression AI**", "Chạy lại golden set mỗi khi đổi prompt hoặc model", "Không tụt so với baseline"],
        ["**Marker eContract**", "EC-01 … EC-09 theo `07-econtract-integration.md`", "Pass toàn bộ; FPT trả `envelopeId`"],
        ["**Bảo mật**", "Thử vượt quyền, thử ghi vùng khoá qua API trực tiếp", "Bị chặn và ghi audit"],
    ], widths=[3.2, 6.4, 6.9])

    callout(doc, "Ca kiểm thử đầu tiên của golden set nên là chính file đã khảo sát: nó chứa đủ mọi ca khó "
                 "(vùng rỗng, vùng 2 ký tự, vùng 3.174 ký tự, vùng bắc qua bảng, comment sẵn có, numbering tự động) "
                 "và một lỗi số ↔ chữ có thật để chứng minh giá trị hệ thống.")

    # ---------------- X ----------------
    h1(doc, "X. RỦI RO")

    table(doc, ["#", "Rủi ro", "Mức", "Giảm thiểu"], [
        ["1", "**Tiến độ** — backend 0 dòng code, mốc pilot tháng 10, scope rộng", "Rất cao",
         "Cắt scope theo QĐ-2; triển khai theo lát cắt dọc; chốt phương án ở mục XI"],
        ["2", "**Write-back vỡ format** ở 4 vùng block / cross-table", "Cao",
         "PoC-1 làm cổng chặn; chế độ C (chỉ cảnh báo) là đường lùi an toàn"],
        ["3", "**AI bỏ sót điều khoản Red Flag**", "Cao",
         "Golden set + ngưỡng recall riêng; disclaimer C-9; Legal vẫn duyệt cuối"],
        ["4", "**Chưa có credentials FPT** chặn kiểm thử tích hợp", "Cao",
         "Adapter + mock server ngay tuần 1; leo thang xin credentials là việc ngày 1"],
        ["5", "**License AGPLv3** nếu về sau nhúng editor", "Trung bình",
         "QĐ-1 đã gỡ AGPL khỏi *đường ghi*; chỉ còn rủi ro ở lớp hiển thị Phase 1.5. Hỏi IT Governance song song, "
         "dự phòng ngân sách commercial"],
    ], widths=[0.9, 5.4, 2.0, 8.2])

    para(doc, "Theo dõi thêm nhưng chưa vào top 5: resolve numbering sai làm lệch trích dẫn điều khoản; "
              "comment orphaned hàng loạt sau PT3; nghẽn queue cuối quý.")

    # ---------------- XI ----------------
    h1(doc, "XI. KẾ HOẠCH TRIỂN KHAI")

    h2(doc, "1. PoC cổng chặn — 2 tuần đầu, chạy song song")
    table(doc, ["PoC", "Nội dung", "Thời lượng", "Tiêu chí đạt"], [
        ["**PoC-1**", "`lxml` round-trip và ghi 3 chế độ trên template thật", "1 tuần",
         "Mở bằng Word không cảnh báo · diff XML ngoài allow-list rỗng · ghi đúng 11 atomic + 2 block · "
         "2 vùng cross-table ít nhất không vỡ bảng"],
        ["**PoC-2**", "Marker mực trắng → FPT demo env", "3 ngày",
         "FPT trả `envelopeId`. Nếu chưa có credentials thì chạy mock và ghi nhận nợ kỹ thuật"],
        ["**PoC-3**", "vLLM + Qwen3-30B-A3B guided JSON trên một checklist thật", "3 ngày",
         "50 lần gọi liên tiếp đều trả JSON hợp lệ schema, p95 ≤ 8 s/clause"],
    ], widths=[1.6, 5.0, 2.0, 7.9])
    para(doc, "**PoC-1 là cổng chặn cứng.** Nếu trượt, thu hẹp Phase 1 còn 11 atomic field, 4 vùng block chuyển "
              "sang chế độ chỉ cảnh báo.")

    h2(doc, "2. Lát cắt dọc")
    table(doc, ["Lát cắt", "Nội dung", "Chứng minh được"], [
        ["VS-1", "Login → tạo review từ template → ingest → hiện inventory 16 vùng mở", "QĐ-8 và tầng OOXML đọc"],
        ["VS-2", "VS-1 + chạy AI review → findings + 2 điểm số → ghi một atomic field", "Trục AI xuyên suốt"],
        ["VS-3", "VS-2 + marker kéo-thả + submit → Manager → Legal → outbox eContract", "Trục nghiệp vụ xuyên suốt"],
        ["VS-4", "VS-3 + PT3 reupload + structural binding + comment TH1", "Đủ cho pilot"],
        ["VS-5", "Editor nhúng + TH2 Track Changes", "Phase 1.5"],
    ], widths=[1.8, 8.2, 6.5])

    h2(doc, "3. Ước tính công sức")
    para(doc, "Đơn vị người-tuần (pw), giả định đội **2 backend/AI + 1 frontend**.")
    table(doc, ["#", "Module", "pw", "#", "Module", "pw"], [
        ["1", "`services/document/` — OOXML ★", "**8**", "6", "eContract outbound + outbox", "3"],
        ["2", "`services/ai/` — pipeline ★", "**7**", "7", "Config checklist (API mới)", "3"],
        ["3", "Data model + migration + audit", "3", "8", "Frontend điều chỉnh", "5"],
        ["4", "API + state machine + RBAC", "4", "9", "Golden set + eval harness", "2"],
        ["5", "Queue + worker + realtime", "2", "10", "Hạ tầng, CI, observability", "2"],
        ["", "", "", "", "**Tổng Phase 1**", "**39 pw**"],
        ["", "", "", "", "*Editor + TH2 (Phase 1.5)*", "*+6*"],
    ], widths=[0.9, 6.0, 1.4, 0.9, 6.0, 1.4])
    para(doc, "Với 3 người: **39 / 3 ≈ 13 tuần lịch**, chưa trừ rủi ro.")

    h2(doc, "4. Đánh giá tính khả thi")
    callout(doc, "Nói thẳng: **mốc test tháng 9 không khả thi.** Còn khoảng 4 tuần, trong đó 2 tuần đầu dành cho PoC. "
                 "Đến mốc pilot tháng 10 còn khoảng 9 tuần ≈ 27 pw, so với 39 pw của Phase 1 rút gọn thì **thiếu "
                 "khoảng 12 pw**.", color=DANGER, fill="FBECEC")
    table(doc, ["Phương án", "Nội dung", "Rủi ro"], [
        ["**PA-1 — khuyến nghị**", "Pilot tháng 10 = VS-3 với **một loại hợp đồng khung duy nhất**, một checklist, "
                                   "eContract chạy môi trường Demo. VS-4 bổ sung tháng 11. Editor và TH2 sang Phase 1.5",
         "Thấp — có thứ chạy thật đúng hạn"],
        ["PA-2", "Giữ nguyên scope, lùi pilot sang tháng 12", "Trung bình — ảnh hưởng niềm tin stakeholder"],
        ["PA-3", "Giữ nguyên scope và mốc, tăng lên 5 người",
         "Cao — người mới khó tăng tốc ở đúng 2 module khó nhất"],
    ], widths=[3.4, 9.1, 4.0])
    para(doc, "Nếu không có quyết định nào được chọn, thiết kế này mặc định chạy theo **PA-1**.")

    # ---------------- XII ----------------
    h1(doc, "XII. VIỆC CẦN CHỐT NGAY")

    table(doc, ["#", "Việc", "Người", "Chặn cái gì"], [
        ["1", "Cung cấp **2–3 template hợp đồng khung** khác để xác nhận PH-1…PH-10 có tổng quát không", "BA + Legal", "Toàn bộ tầng tài liệu"],
        ["2", "Xác nhận **có bản template trắng** do Legal ban hành không, hay Purchasing luôn copy hợp đồng cũ", "Legal", "Đường instantiate (QĐ-8)"],
        ["3", "Leo thang xin **credentials FPT môi trường Demo**, `selector`, `docTypeCode`", "PM", "PoC-2, kiểm thử EC"],
        ["4", "Xác nhận **A100 40GB hay 80GB**", "IT", "Cấu hình vLLM, model artifact"],
        ["5", "Chốt **PA-1 / PA-2 / PA-3** ở mục XI.4", "PM", "Toàn bộ kế hoạch"],
        ["6", "Xác nhận **số clause mỗi loại HĐ** và cam kết nguồn lực gán nhãn golden set", "Legal", "Capacity và nghiệm thu AI"],
        ["7", "Bổ sung Blueprint: **đường escalate khi người duyệt yêu cầu sửa vùng khoá**", "BA", "Luồng TH1/TH2"],
        ["8", "Hỏi **IT Governance về chính sách AGPLv3**", "PM", "Chỉ chặn Phase 1.5"],
    ], widths=[0.9, 8.2, 2.4, 5.0])

    # ---------------- XIII ----------------
    h1(doc, "XIII. NGOÀI PHẠM VI PHASE 1")
    bullets(doc, [
        "Chiều nhận file đã ký từ eContract — do hệ thống hiện hữu đảm nhận (C-5).",
        "Hợp đồng nhà cung cấp, đặc biệt trường hợp không có vùng mở nào — Phase 2.",
        "Track Changes trên UI (TH2) — hoãn sang Phase 1.5 theo QĐ-2.",
        "Import/Export checklist bằng file — đã bỏ khỏi Sprint 1.",
        "Auto-routing theo Approval Matrix — Sprint 1 chỉ cảnh báo (C-6).",
        "Fine-tune model chuyên ngành luật.",
        "Legacy Form Field — template thật không dùng; chỉ giữ khả năng phát hiện và cảnh báo.",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Hết —")
    r.italic = True
    r.font.color.rgb = MUTED

    doc.save(path)
    print(f"Đã tạo: {path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else \
        "docs/AI-Legal_Technical-Solution_Phase1.docx"
    build(out)
