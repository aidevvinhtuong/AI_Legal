#!/usr/bin/env python3
"""
md-to-docx.py — gộp toàn bộ docs/technical-solution/TS-*.md thành MỘT file .docx.

Đặc điểm:
  - Giữ đúng cấu trúc heading, bảng, code block, blockquote, danh sách
  - Sơ đồ ```mermaid được RENDER THÀNH ẢNH bằng mermaid-cli (nếu có), không dán code
  - Trang bìa + mục lục tự động (Word cập nhật khi mở) + header/footer đánh số trang
  - Bỏ qua README.md

Phụ thuộc: python-docx (đã có sẵn). mermaid-cli là tuỳ chọn — thiếu thì sơ đồ
hiển thị dưới dạng khối code có khung.

    python3 scripts/md-to-docx.py [output.docx]
"""
from __future__ import annotations

import hashlib
import os
import re
import struct
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# Cấu hình
# ─────────────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "docs" / "technical-solution"
DEFAULT_OUT = ROOT / "docs" / "AI-Legal_Technical-Solution_Phase1_Full.docx"
CACHE_DIR = Path(os.environ.get("MMD_CACHE", "/tmp/ailegal-mmd-cache"))

MMDC = Path("/tmp/mmdc/node_modules/.bin/mmdc")
MMD_CONFIG = ROOT / "docs" / "diagrams" / "mermaid-config.json"
MMD_PUPPETEER = Path("/tmp/mmdc/puppeteer.json")

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x5A, 0x6B, 0x7C)
CODE_BG = "F4F6F8"
HEAD_BG = "EAF1F8"
QUOTE_BG = "FFF8E6"

CONTENT_W = Cm(17.0)        # A4 21cm − lề 2×2cm
MAX_IMG_H = Cm(20.0)

# Emoji ưu tiên dùng làm nhãn trong bảng — Word hay render thành ô vuông.
EMOJI_STRIP = {"🔴": "", "🟠": "", "🟡": "", "🟢": "", "⚪": "", "★": "*", "👤": "", "🏛": ""}


# ─────────────────────────────────────────────────────────────────────────────
# Tiện ích OOXML
# ─────────────────────────────────────────────────────────────────────────────
def shade(el, fill: str) -> None:
    """Tô nền cho paragraph (pPr) hoặc cell (tcPr)."""
    pr = el.get_or_add_pPr() if hasattr(el, "get_or_add_pPr") else el
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def cell_shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def left_bar(par, color: str = "1F4E79") -> None:
    """Vạch dọc bên trái — dùng cho code block và blockquote."""
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    bdr.append(left)
    pPr.append(bdr)


def keep_with_next(par) -> None:
    pPr = par._p.get_or_add_pPr()
    k = OxmlElement("w:keepNext")
    pPr.append(k)


def repeat_header(row) -> None:
    """Lặp lại hàng tiêu đề khi bảng tràn sang trang mới."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def add_field(par, instr: str) -> None:
    """Chèn field code Word (TOC, PAGE, NUMPAGES)."""
    r1 = par.add_run()._r
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r1.append(fc)
    r2 = par.add_run()._r
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2.append(it)
    r3 = par.add_run()._r
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r3.append(fs)
    r4 = par.add_run("…")._r
    r5 = par.add_run()._r
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r5.append(fe)


def add_hyperlink(par, text: str, url: str):
    part = par.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1F4E79"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    par._p.append(link)


# ─────────────────────────────────────────────────────────────────────────────
# Style
# ─────────────────────────────────────────────────────────────────────────────
def setup_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    sizes = {"Heading 1": 18, "Heading 2": 14.5, "Heading 3": 12.5, "Heading 4": 11.5}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(14 if size > 13 else 10)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21), Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2)
        sec.top_margin = sec.bottom_margin = Cm(2)


def add_footer(section, label: str) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{label}    |    Trang ")
    r.font.size = Pt(8.5); r.font.color.rgb = MUTED; r.font.name = BODY_FONT
    add_field(p, "PAGE")
    r2 = p.add_run(" / ")
    r2.font.size = Pt(8.5); r2.font.color.rgb = MUTED
    add_field(p, "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(8.5); run.font.color.rgb = MUTED; run.font.name = BODY_FONT


# ─────────────────────────────────────────────────────────────────────────────
# Inline markdown
# ─────────────────────────────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r"(`[^`]+`)"                      # code
    r"|(\*\*\*[^*]+\*\*\*)"           # bold italic
    r"|(\*\*[^*]+?\*\*)"              # bold
    r"|(\*[^*\s][^*]*?\*)"            # italic
    r"|(\[[^\]]+\]\([^)]+\))"         # link
)


def clean_text(s: str) -> str:
    for k, v in EMOJI_STRIP.items():
        s = s.replace(k, v)
    return s.replace("\\|", "|").replace("\\_", "_").replace("\\*", "*")


def add_inline(par, text: str, base_size: float = 10.5, base_color=None) -> None:
    """Render markdown inline vào một paragraph."""
    text = clean_text(text)
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            _plain(par, text[pos:m.start()], base_size, base_color)
        tok = m.group(0)
        if tok.startswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = RGBColor(0xA6, 0x1B, 0x1B)
        elif tok.startswith("***"):
            r = _plain(par, tok[3:-3], base_size, base_color); r.bold = True; r.italic = True
        elif tok.startswith("**"):
            r = _plain(par, tok[2:-2], base_size, base_color); r.bold = True
        elif tok.startswith("*"):
            r = _plain(par, tok[1:-1], base_size, base_color); r.italic = True
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok).groups()
            if url.startswith(("http://", "https://")):
                add_hyperlink(par, label, url)
            else:
                # Link nội bộ giữa các file .md — trong docx gộp thì vô nghĩa,
                # chỉ giữ nhãn.
                r = _plain(par, label, base_size, base_color)
                r.font.color.rgb = ACCENT
        pos = m.end()
    if pos < len(text):
        _plain(par, text[pos:], base_size, base_color)


def _plain(par, s: str, size: float, color):
    r = par.add_run(s)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


# ─────────────────────────────────────────────────────────────────────────────
# Mermaid
# ─────────────────────────────────────────────────────────────────────────────
def png_size(path: Path) -> tuple[int, int]:
    """Đọc kích thước PNG từ chunk IHDR — không cần thư viện ảnh."""
    with path.open("rb") as f:
        head = f.read(24)
    if head[:8] != b"\x89PNG\r\n\x1a\n":
        raise ValueError("không phải PNG")
    return struct.unpack(">II", head[16:24])


def render_mermaid(code: str) -> Path | None:
    """Render mermaid → PNG. Có cache theo hash nội dung."""
    if not MMDC.exists():
        return None
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    h = hashlib.sha256(code.encode()).hexdigest()[:16]
    out = CACHE_DIR / f"{h}.png"
    if out.exists():
        return out

    src = CACHE_DIR / f"{h}.mmd"
    src.write_text(code, encoding="utf-8")
    cmd = [str(MMDC), "-i", str(src), "-o", str(out), "-b", "white", "-s", "2"]
    if MMD_CONFIG.exists():
        cmd += ["-c", str(MMD_CONFIG)]
    if MMD_PUPPETEER.exists():
        cmd += ["-p", str(MMD_PUPPETEER)]
    env = {**os.environ, "PUPPETEER_CACHE_DIR": "/tmp/mmdc/.puppeteer"}
    try:
        r = subprocess.run(cmd, capture_output=True, timeout=120, env=env)
    except subprocess.TimeoutExpired:
        print("      ! mermaid timeout", file=sys.stderr)
        return None
    if r.returncode != 0 or not out.exists():
        print(f"      ! mermaid lỗi: {r.stderr.decode()[:200]}", file=sys.stderr)
        return None
    return out


def add_diagram(doc: Document, code: str, caption: str) -> bool:
    png = render_mermaid(code)
    if png is None:
        return False
    w_px, h_px = png_size(png)
    width = CONTENT_W
    height = int(width * h_px / w_px)
    if height > MAX_IMG_H:
        height = MAX_IMG_H
        width = int(height * w_px / h_px)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(png), width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = MUTED; r.font.name = BODY_FONT
    return True


# ─────────────────────────────────────────────────────────────────────────────
# Khối nội dung
# ─────────────────────────────────────────────────────────────────────────────
def add_code_block(doc: Document, lines: list[str], lang: str) -> None:
    if lang:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(0)
        keep_with_next(h)
        r = h.add_run(lang.upper())
        r.font.size = Pt(7.5); r.font.bold = True; r.font.color.rgb = MUTED; r.font.name = BODY_FONT

    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(3 if i == 0 else 0)
        pf.space_after = Pt(3 if i == len(lines) - 1 else 0)
        pf.left_indent = Cm(0.3)
        pf.line_spacing = 1.0
        shade(p._p, CODE_BG)
        left_bar(p, "C3CCD5")
        r = p.add_run(line if line.strip() else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(8)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)


def split_row(line: str) -> list[str]:
    """Tách ô của bảng markdown, tôn trọng dấu \\| escape."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    out, cur, i = [], "", 0
    while i < len(line):
        if line[i] == "\\" and i + 1 < len(line) and line[i + 1] == "|":
            cur += "\\|"; i += 2; continue
        if line[i] == "|":
            out.append(cur); cur = ""; i += 1; continue
        cur += line[i]; i += 1
    out.append(cur)
    return [c.strip() for c in out]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]

    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True

    for ri, row in enumerate(rows):
        for ci, text in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, text, base_size=9)
            if ri == 0:
                cell_shade(cell, HEAD_BG)
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = ACCENT
    repeat_header(t.rows[0])

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(8)
    after.paragraph_format.space_before = Pt(0)


def add_quote(doc: Document, lines: list[str]) -> None:
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 0)
        shade(p._p, QUOTE_BG)
        left_bar(p, "B8860B")
        add_inline(p, line, base_size=10)


# ─────────────────────────────────────────────────────────────────────────────
# Parser
# ─────────────────────────────────────────────────────────────────────────────
FENCE_RE = re.compile(r"^```(\w*)\s*$")
HEAD_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ULI_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
OLI_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
HR_RE = re.compile(r"^\s*([-*_])\s*(\1\s*){2,}$")
TBL_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|[\s:|-]*$")


def render_markdown(doc: Document, text: str, heading_offset: int, doc_label: str) -> int:
    """Đổ nội dung markdown vào document. Trả về số sơ đồ đã render."""
    lines = text.split("\n")
    i, n = 0, len(lines)
    diagrams = 0
    last_heading = doc_label

    while i < n:
        line = lines[i]

        # ── code fence ────────────────────────────────────────────────────
        m = FENCE_RE.match(line)
        if m:
            lang = m.group(1).lower()
            body, i = [], i + 1
            while i < n and not FENCE_RE.match(lines[i]):
                body.append(lines[i]); i += 1
            i += 1
            if lang == "mermaid":
                diagrams += 1
                cap = f"Sơ đồ {doc_label}.{diagrams} — {last_heading}"
                if not add_diagram(doc, "\n".join(body), cap):
                    add_code_block(doc, body, "mermaid")
            else:
                add_code_block(doc, body, lang)
            continue

        # ── heading ───────────────────────────────────────────────────────
        m = HEAD_RE.match(line)
        if m:
            level = min(len(m.group(1)) + heading_offset, 9)
            title = clean_text(m.group(2)).strip()
            # Bỏ số thứ tự và ký hiệu markdown còn sót để caption sơ đồ sạch
            last_heading = re.sub(r"^[IVXLC]+\.\s*|^\d+(\.\d+)*\.?\s*", "", title)
            last_heading = last_heading.replace("`", "").replace("**", "").strip()[:60]
            h = doc.add_heading(level=level)
            add_inline(h, m.group(2), base_size={1: 18, 2: 14.5, 3: 12.5}.get(level, 11.5))
            for r in h.runs:
                r.font.bold = True
                r.font.color.rgb = ACCENT
            i += 1
            continue

        # ── bảng ──────────────────────────────────────────────────────────
        if line.lstrip().startswith("|") and i + 1 < n and TBL_SEP_RE.match(lines[i + 1]):
            rows = [split_row(line)]
            i += 2
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            add_table(doc, rows)
            continue

        # ── blockquote ────────────────────────────────────────────────────
        if line.lstrip().startswith(">"):
            body = []
            while i < n and lines[i].lstrip().startswith(">"):
                body.append(lines[i].lstrip()[1:].strip()); i += 1
            add_quote(doc, [b for b in body if b] or [""])
            continue

        # ── đường kẻ ngang ────────────────────────────────────────────────
        if HR_RE.match(line):
            i += 1
            continue

        # ── danh sách ─────────────────────────────────────────────────────
        m = ULI_RE.match(line)
        if m:
            depth = min(len(m.group(1)) // 2, 2)
            p = doc.add_paragraph(style="List Bullet" if depth == 0 else "List Bullet 2")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.6 + 0.5 * depth)
            add_inline(p, m.group(2))
            i += 1
            continue

        m = OLI_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, m.group(3))
            i += 1
            continue

        # ── đoạn văn ──────────────────────────────────────────────────────
        if line.strip():
            buf = []
            while i < n and lines[i].strip() and not (
                HEAD_RE.match(lines[i]) or FENCE_RE.match(lines[i])
                or lines[i].lstrip().startswith((">", "|"))
                or ULI_RE.match(lines[i]) or OLI_RE.match(lines[i])
                or HR_RE.match(lines[i])
            ):
                buf.append(lines[i].strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, " ".join(buf))
            continue

        i += 1

    return diagrams


# ─────────────────────────────────────────────────────────────────────────────
# Trang bìa & mục lục
# ─────────────────────────────────────────────────────────────────────────────
def build_cover(doc: Document, files: list[Path]) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TECHNICAL SOLUTION")
    r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = ACCENT; r.font.name = BODY_FONT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 1 — Hợp đồng khung")
    r.font.size = Pt(16); r.font.color.rgb = MUTED; r.font.name = BODY_FONT

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI LEGAL — Hệ thống AI Review Hợp đồng")
    r.font.size = Pt(15); r.font.bold = True; r.font.name = BODY_FONT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Saint-Gobain Việt Nam")
    r.font.size = Pt(12); r.font.color.rgb = MUTED; r.font.name = BODY_FONT

    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta = [
        ("Phiên bản", "1.0"),
        ("Ngày phát hành", date.today().strftime("%d/%m/%Y")),
        ("Người thực hiện", "Backend / AI Engineer"),
        ("Số tài liệu thành phần", f"{len(files)} (TS-00 → TS-11)"),
    ]
    for ri, (k, v) in enumerate(meta):
        for ci, txt in enumerate((k, v)):
            cell = t.cell(ri, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(10); run.font.name = BODY_FONT
            if ci == 0:
                run.font.bold = True; run.font.color.rgb = ACCENT
                cell_shade(cell, HEAD_BG)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu nội bộ — không phổ biến ra ngoài công ty")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = MUTED; r.font.name = BODY_FONT

    doc.add_page_break()


def build_toc(doc: Document) -> None:
    h = doc.add_heading(level=1)
    r = h.add_run("MỤC LỤC")
    r.font.bold = True; r.font.color.rgb = ACCENT; r.font.size = Pt(18); r.font.name = BODY_FONT

    note = doc.add_paragraph()
    r = note.add_run(
        "Mục lục được sinh tự động. Khi mở bằng Microsoft Word: bấm Ctrl+A rồi F9 "
        "(hoặc chuột phải vào mục lục → Update Field → Update entire table) để cập nhật số trang."
    )
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = MUTED; r.font.name = BODY_FONT

    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────
def main(argv: list[str]) -> int:
    out_path = Path(argv[0]).resolve() if argv else DEFAULT_OUT

    files = sorted(p for p in SRC_DIR.glob("TS-*.md"))
    if not files:
        print(f"Không tìm thấy TS-*.md trong {SRC_DIR}", file=sys.stderr)
        return 1

    print(f"Nguồn : {SRC_DIR} ({len(files)} file, đã bỏ README.md)")
    print(f"Mermaid: {'có mermaid-cli — render ảnh' if MMDC.exists() else 'KHÔNG có — sơ đồ giữ dạng code'}")

    doc = Document()
    setup_document(doc)
    build_cover(doc, files)
    build_toc(doc)
    add_footer(doc.sections[0], "AI Legal · Technical Solution Phase 1 · v1.0")

    total_diagrams = 0
    for idx, f in enumerate(files):
        label = f.stem.split("-")[1]            # '00', '01', …
        print(f"  [{idx + 1}/{len(files)}] {f.name}")
        if idx > 0:
            doc.add_page_break()
        text = f.read_text(encoding="utf-8")
        total_diagrams += render_markdown(doc, text, heading_offset=0, doc_label=label)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    size_kb = out_path.stat().st_size / 1024
    print(f"\nĐã ghi: {out_path}")
    print(f"        {size_kb:,.0f} KB · {total_diagrams} sơ đồ đã render")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
